"""
File Handler Service
=====================
Orchestrates the processing pipeline based on file type.
Routes images, PDFs, and DOCX files through the appropriate
preprocessing → OCR → layout → extraction chain.
"""

import cv2
import numpy as np
from pathlib import Path
from PIL import Image
import fitz  # PyMuPDF
from docx import Document
from services.preprocessing import preprocess_image
from services.ocr_service import run_ocr
from services.layout_engine import analyze_layout
from services.extractor import extract_invoice_data
from services.validator import validate_extraction
from services.json_builder import build_json_response
from services.table_extractor import extract_tables
from pdfminer.high_level import extract_text as pdfminer_extract_text


def process_file(
    file_path: str,
    file_ext: str,
    file_id: str,
    enable_validation: bool = False,
) -> dict:
    """
    Main processing orchestrator.
    Routes the file through the appropriate pipeline based on its type.

    Returns a structured JSON-serializable dict.
    """
    if file_ext in (".jpg", ".jpeg", ".png"):
        ocr_results = _process_image(file_path)
    elif file_ext == ".pdf":
        ocr_results = _process_pdf(file_path)
    elif file_ext == ".docx":
        ocr_results = _process_docx(file_path)
    else:
        raise ValueError(f"Unsupported file extension: {file_ext}")

    # Merge all OCR results into a single dataset
    merged = _merge_ocr_results(ocr_results)

    # Layout analysis
    layout = analyze_layout(merged)

    # Data extraction
    extracted = extract_invoice_data(merged, layout)

    # Extract runtime metadata injected by orchestrator
    template_used = extracted.pop("_runtime_template_used_", "generic")

    # Extract Native Tables
    tables = []
    if file_ext == ".pdf":
        tables = extract_tables(file_path)

    # Optional validation
    validation = None
    if enable_validation:
        validation = validate_extraction(extracted, merged)

    # Build final JSON
    return build_json_response(
        file_id=file_id,
        extracted=extracted,
        raw_text=merged["raw_text"],
        validation=validation,
        engine=merged.get("engine", "tesseract"),
        template=template_used,
        tables_detected=len(tables) > 0,
        line_items=tables
    )



def _process_image(file_path: str) -> list:
    """Process a single image file."""
    img = cv2.imread(file_path)
    if img is None:
        raise ValueError(f"Could not read image: {file_path}")

    preprocessed = preprocess_image(img)
    ocr_result = run_ocr(preprocessed)
    return [ocr_result]


def _process_pdf(file_path: str) -> list:
    """
    Dual-Pass PDF Extraction:
    Pass 1: Use `pdfplumber` for mathematically perfect native digital data extraction.
    Pass 2: Use `PyMuPDF` + OpenCV + OCR for scanned image pages fallback.
    """
    import pdfplumber
    results = []
    
    # ── Pass 1: Try Native Digital Extraction ─────────────────────
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                raw_text = page.extract_text()
                try:
                    plumber_words = page.extract_words()
                except Exception:
                    plumber_words = []
                
                # If there's plenty of digital text, use it natively without OCR!
                if raw_text and len(plumber_words) > 10:
                    words = []
                    zoom = 300 / 72  # Upscale coordinates to match 300 DPI OCR bounds
                    for w in plumber_words:
                        x = w['x0'] * zoom
                        y = w['top'] * zoom
                        width = (w['x1'] - w['x0']) * zoom
                        height = (w['bottom'] - w['top']) * zoom
                        
                        words.append({
                            "text": w['text'],
                            "x": x,
                            "y": y,
                            "w": width,
                            "h": height,
                            "conf": 100,  # Native format is structurally 100% accurate
                            "block_num": 1,
                            "line_num": page_num, # Fallback line approximation
                            "word_num": len(words) + 1,
                        })
                    
                    results.append({
                        "words": words,
                        "raw_text": raw_text,
                        "engine": "pdfplumber",
                        "confidence": 100
                    })
                else:
                    # Tier 2: Try PyMuPDF (fitz) native text extraction
                    fitz_doc = fitz.open(file_path)
                    fitz_page = fitz_doc[page_num - 1]
                    fitz_text = fitz_page.get_text("text").strip()
                    fitz_doc.close()
                    
                    if len(fitz_text) > 20: # Has text, wasn't caught by pdfplumber
                        results.append({
                            "words": _text_to_synthetic_words(fitz_text),
                            "raw_text": fitz_text,
                            "engine": "pymupdf",
                            "confidence": 95
                        })
                    else:
                        # Tier 3: Try pdfminer.six text fallback
                        miner_text = pdfminer_extract_text(file_path, page_numbers=[page_num - 1]).strip()
                        if len(miner_text) > 20:
                            results.append({
                                "words": _text_to_synthetic_words(miner_text),
                                "raw_text": miner_text,
                                "engine": "pdfminer.six",
                                "confidence": 90
                            })
                        else:
                            # ALL native text extractions failed. The PDF is guaranteed to be a scanned image.
                            # Append None to trigger OCR pass 4 targeted exclusively on this page
                            results.append(None)
    except Exception as e:
        print(f"[File Handler] pdfplumber native extraction failed: {e}. Falling back to full OCR.")
        results = [None] # Force full OCR fallback

    # ── Pass 2: Fallback OCR for scanned pages ────────────────────
    if None in results:
        try:
            doc = fitz.open(file_path)
            zoom = 300 / 72  # 300 DPI
            mat = fitz.Matrix(zoom, zoom)
            
            for i in range(len(doc)):
                # If we parsed this particular page natively, skip OCR!
                if len(results) > i and results[i] is not None:
                    continue
                    
                page = doc[i]
                pix = page.get_pixmap(matrix=mat, alpha=False)
                img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, pix.n)
                
                if pix.n == 3:
                    img_bgr = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
                elif pix.n == 1:
                    img_bgr = cv2.cvtColor(img_array, cv2.COLOR_GRAY2BGR)
                else:
                    img_bgr = img_array

                preprocessed = preprocess_image(img_bgr)
                ocr_result = run_ocr(preprocessed)
                
                # Replace the None placeholder with physical OCR results
                if len(results) <= i:
                    results.append(ocr_result)
                else:
                    results[i] = ocr_result
                    
            doc.close()
        except Exception as e:
            raise ValueError(f"PDF fallback OCR failed: {str(e)}")

    # Clean up and validate
    final_results = [r for r in results if r is not None]
    if not final_results:
        raise ValueError("No text or processable layout found in PDF file.")
        
    return final_results


def _process_docx(file_path: str) -> list:
    """Extract text and embedded images from DOCX, OCR the images."""
    doc = Document(file_path)
    results = []

    # Extract body text directly (no OCR needed)
    body_text_parts = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            body_text_parts.append(text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                body_text_parts.append(" | ".join(row_texts))

    body_text = "\n".join(body_text_parts)

    # Create a synthetic OCR result for the text content
    if body_text:
        text_result = {
            "words": _text_to_synthetic_words(body_text),
            "raw_text": body_text,
        }
        results.append(text_result)

    # Extract embedded images and run OCR on them
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                image_data = rel.target_part.blob
                img_array = np.frombuffer(image_data, dtype=np.uint8)
                img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

                if img is not None:
                    preprocessed = preprocess_image(img)
                    ocr_result = run_ocr(preprocessed)
                    results.append(ocr_result)
            except Exception:
                # Skip unreadable embedded images
                continue

    if not results:
        raise ValueError("No text or processable images found in DOCX file.")

    return results


def _text_to_synthetic_words(text: str) -> list:
    """
    Convert plain text into synthetic 'word' dicts that mimic OCR output.
    These won't have real bounding boxes but provide the text for extraction.
    """
    words = []
    y_offset = 0
    for line_num, line in enumerate(text.split("\n"), start=1):
        x_offset = 0
        for word_num, word in enumerate(line.split(), start=1):
            words.append({
                "text": word,
                "x": x_offset,
                "y": y_offset,
                "w": len(word) * 10,
                "h": 20,
                "conf": 100,  # Text extraction is 100% confident
                "block_num": 1,
                "line_num": line_num,
                "word_num": word_num,
            })
            x_offset += len(word) * 10 + 10
        y_offset += 30
    return words


def _merge_ocr_results(results: list) -> dict:
    """Merge multiple OCR result dicts into a single unified result."""
    all_words = []
    raw_text_parts = []
    engines_used = set()

    y_page_offset = 0
    for result in results:
        for word in result.get("words", []):
            # Offset y-coordinates for multi-page merging
            merged_word = word.copy()
            merged_word["y"] += y_page_offset
            all_words.append(merged_word)

        raw_text_parts.append(result.get("raw_text", ""))
        engines_used.add(result.get("engine", "tesseract"))

        # Calculate page height for offset (use max y + h)
        if result.get("words"):
            max_y = max(w["y"] + w["h"] for w in result["words"])
            y_page_offset += max_y + 50  # 50px gap between pages

    engine_val = "hybrid" if len(engines_used) > 1 else (engines_used.pop() if engines_used else "tesseract")

    return {
        "words": all_words,
        "raw_text": "\n\n--- PAGE BREAK ---\n\n".join(raw_text_parts),
        "engine": engine_val
    }
